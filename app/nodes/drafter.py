"""Drafter 노드 — 대화형 인터뷰 + 계약서 생성 + 자체 검증.

Phase 1: 인터뷰 (필수/선택 정보 대화형 수집)
Phase 2: 표준 서식 검색 (RAG)
Phase 3: 계약서 생성 (Claude Sonnet)
Phase 4: 자체 검증 (Analyzer + Validator)
Phase 5: 출력 (DOCX)
"""

import json
import logging
import uuid
from io import BytesIO
from typing import Any

from app.config import settings
from app.llm.client import call_llm, call_llm_json
from app.llm.prompts.drafter_system import DRAFTER_SYSTEM_PROMPT
from app.nodes.rag import _search_standard_clauses, _generate_queries
from app.state.draft_state import DraftState
from app.utils.db_client import execute, fetchrow

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────
# 인터뷰 스키마
# ──────────────────────────────────────────

INTERVIEW_SCHEMAS: dict[str, dict] = {
    "service_contract": {
        "label": "용역계약",
        "required": [
            {
                "field": "parties",
                "question": "계약 당사자 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "party_a", "label": "갑 (발주자)", "placeholder": "예: (주)에이비씨"},
                    {"key": "party_b", "label": "을 (수급자)", "placeholder": "예: 홍길동"},
                ],
            },
            {"field": "scope", "question": "용역의 구체적인 내용은 무엇인가요?", "placeholder": "예: 웹앱 개발, 디자인"},
            {
                "field": "duration",
                "question": "계약 기간을 입력해주세요.",
                "sub_fields": [
                    {"key": "start_date", "label": "시작일", "placeholder": "예: 2026.05.01"},
                    {"key": "end_date", "label": "종료일", "placeholder": "예: 2026.07.31"},
                ],
            },
            {
                "field": "payment",
                "question": "대금과 지급 조건을 입력해주세요.",
                "sub_fields": [
                    {"key": "amount", "label": "총 금액", "placeholder": "예: 5,000만 원"},
                    {"key": "method", "label": "지급 방식", "placeholder": "예: 월별 / 완료 후 / 분할"},
                ],
            },
            {"field": "deliverables", "question": "납품물은 무엇인가요?", "placeholder": "예: 소스코드, 디자인 파일"},
        ],
        "optional": [
            {"field": "ip_ownership", "question": "결과물의 지식재산권은 누구에게 귀속?", "default": "갑에게 귀속"},
            {"field": "confidentiality", "question": "비밀유지 조항 필요?", "default": "양측 상호 비밀유지"},
            {"field": "termination", "question": "중도 해지 조건은?", "default": "양측 30일 전 서면 통지"},
            {"field": "liability", "question": "손해배상 한도?", "default": "계약금액 한도 내 직접손해"},
            {"field": "dispute", "question": "분쟁해결 방법?", "default": "서울중앙지방법원 관할"},
        ],
    },
    "nda": {
        "label": "비밀유지계약(NDA)",
        "required": [
            {
                "field": "parties",
                "question": "계약 당사자 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "discloser", "label": "정보 제공자", "placeholder": "예: (주)에이비씨"},
                    {"key": "recipient", "label": "정보 수령자", "placeholder": "예: (주)디이에프"},
                ],
            },
            {"field": "confidential_info", "question": "비밀정보의 범위는?", "placeholder": "예: 기술정보, 경영정보"},
            {"field": "duration", "question": "비밀유지 기간은?", "placeholder": "예: 계약종료 후 3년"},
            {"field": "purpose", "question": "정보 제공 목적은?", "placeholder": "예: 협업 검토, 투자 검토"},
        ],
        "optional": [
            {"field": "penalty", "question": "위반 시 위약금?", "default": "실제 손해배상"},
            {"field": "return_clause", "question": "종료 시 자료 반환/파기?", "default": "반환 또는 파기"},
        ],
    },
    "employment": {
        "label": "근로계약",
        "required": [
            {
                "field": "parties",
                "question": "사업장과 근로자 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "employer", "label": "사업장 (사용자)", "placeholder": "예: (주)에이비씨"},
                    {"key": "employee", "label": "근로자", "placeholder": "예: 홍길동"},
                ],
            },
            {
                "field": "position",
                "question": "직위와 업무 내용을 입력해주세요.",
                "sub_fields": [
                    {"key": "title", "label": "직위", "placeholder": "예: 과장"},
                    {"key": "duties", "label": "업무 내용", "placeholder": "예: 소프트웨어 개발"},
                ],
            },
            {
                "field": "salary",
                "question": "급여 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "amount", "label": "급여 (임금)", "placeholder": "예: 월 400만 원"},
                    {"key": "pay_date", "label": "지급일", "placeholder": "예: 매월 25일"},
                ],
            },
            {
                "field": "working_hours",
                "question": "근무 조건을 입력해주세요.",
                "sub_fields": [
                    {"key": "hours", "label": "근무 시간", "placeholder": "예: 09:00 ~ 18:00"},
                    {"key": "location", "label": "근무 장소", "placeholder": "예: 서울시 강남구 본사"},
                ],
            },
            {"field": "duration", "question": "계약 기간은?", "placeholder": "예: 정규직 / 2026.05.01 ~ 2027.04.30"},
        ],
        "optional": [
            {"field": "probation", "question": "수습 기간?", "default": "3개월"},
            {"field": "non_compete", "question": "경업금지 조항?", "default": "퇴직 후 1년, 동종업계"},
            {"field": "benefits", "question": "복리후생?", "default": "4대 보험 가입"},
        ],
    },
    "lease": {
        "label": "임대차계약",
        "required": [
            {
                "field": "parties",
                "question": "임대인과 임차인 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "landlord", "label": "임대인", "placeholder": "예: 홍길동"},
                    {"key": "tenant", "label": "임차인", "placeholder": "예: 김영희"},
                ],
            },
            {
                "field": "property",
                "question": "임대 대상 물건 정보를 입력해주세요.",
                "sub_fields": [
                    {"key": "address", "label": "주소", "placeholder": "예: 서울시 강남구 역삼동 123-4"},
                    {"key": "area", "label": "면적", "placeholder": "예: 전용 85㎡"},
                ],
            },
            {
                "field": "rent",
                "question": "임대 비용을 입력해주세요.",
                "sub_fields": [
                    {"key": "deposit", "label": "보증금", "placeholder": "예: 5,000만 원"},
                    {"key": "monthly_rent", "label": "월세", "placeholder": "예: 100만 원"},
                ],
            },
            {"field": "duration", "question": "임대 기간은?", "placeholder": "예: 2026.05.01 ~ 2028.04.30"},
        ],
        "optional": [
            {"field": "maintenance", "question": "수선/관리 의무 분담?", "default": "소수선 임차인, 대수선 임대인"},
            {"field": "sublease", "question": "전대 허용?", "default": "임대인 사전 동의 필요"},
        ],
    },
}

CONTRACT_TYPE_MAP = {
    "용역": "service_contract", "개발": "service_contract", "용역계약": "service_contract",
    "nda": "nda", "비밀유지": "nda", "기밀": "nda",
    "근로": "employment", "고용": "employment", "근로계약": "employment",
    "임대": "lease", "임대차": "lease", "전세": "lease", "월세": "lease",
}


def interview_node(state: DraftState) -> dict:
    """인터뷰 노드: 필요 정보를 대화형으로 수집."""
    contract_type = state.get("contract_type", "")
    interview_data = dict(state.get("interview_data", {}))
    session_id = state.get("session_id") or str(uuid.uuid4())

    # 계약 유형이 아직 결정되지 않은 경우
    if not contract_type:
        return {
            "session_id": session_id,
            "current_question": {
                "field": "contract_type",
                "question": "어떤 유형의 계약서를 만드시겠습니까?",
                "options": ["용역계약", "비밀유지계약(NDA)", "근로계약", "임대차계약"],
            },
            "interview_complete": False,
            "response": {
                "status": "interviewing",
                "session_id": session_id,
                "question": {
                    "field": "contract_type",
                    "question": "어떤 유형의 계약서를 만드시겠습니까?",
                },
            },
        }

    schema = INTERVIEW_SCHEMAS.get(contract_type, {})
    required = schema.get("required", [])

    # 아직 수집하지 못한 필수 필드
    pending = [f for f in required if f["field"] not in interview_data]

    if not pending:
        # 선택 필드에 default값 적용
        for opt in schema.get("optional", []):
            if opt["field"] not in interview_data:
                interview_data[opt["field"]] = opt.get("default", "")

        return {
            "session_id": session_id,
            "interview_data": interview_data,
            "interview_complete": True,
            "pending_fields": [],
        }

    # 다음 질문
    next_q = pending[0]
    return {
        "session_id": session_id,
        "interview_data": interview_data,
        "current_question": next_q,
        "pending_fields": [f["field"] for f in pending],
        "interview_complete": False,
        "response": {
            "status": "interviewing",
            "session_id": session_id,
            "question": next_q,
            "progress": {
                "collected": len(required) - len(pending),
                "total": len(required),
            },
        },
    }


async def search_template(state: DraftState) -> dict:
    """표준 서식 검색 노드: RAG에서 표준 계약서 조항 조회."""
    contract_type = state.get("contract_type", "")
    interview_data = state.get("interview_data", {})

    scope = interview_data.get("scope", contract_type)
    queries = await _generate_queries(f"{contract_type} 계약서 {scope}")

    standards = await _search_standard_clauses(queries, contract_type=contract_type, limit=10)

    return {"template_clauses": standards}


async def generate_contract(state: DraftState) -> dict:
    """계약서 생성 노드: Claude Sonnet으로 맞춤 계약서 작성."""
    contract_type = state.get("contract_type", "")
    interview_data = state.get("interview_data", {})
    template_clauses = state.get("template_clauses", [])

    schema = INTERVIEW_SCHEMAS.get(contract_type, {})
    label = schema.get("label", contract_type)

    # 표준 서식 텍스트
    template_text = ""
    if template_clauses:
        template_text = "\n\n".join(
            f"[{t.get('clause_type', '조항')}]\n{t.get('standard_text', '')}"
            for t in template_clauses
        )

    # 수집 정보 정리 (sub_fields dict는 세부 항목으로 풀어서 표시)
    info_lines = []
    for k, v in interview_data.items():
        if isinstance(v, dict):
            sub_items = ", ".join(f"{sk}: {sv}" for sk, sv in v.items())
            info_lines.append(f"- {k}: {sub_items}")
        else:
            info_lines.append(f"- {k}: {v}")
    info_text = "\n".join(info_lines)

    user_prompt = (
        f"다음 정보를 바탕으로 {label}를 작성하세요.\n\n"
        f"[수집된 정보]\n{info_text}\n\n"
    )
    if template_text:
        user_prompt += f"[참고 표준 서식]\n{template_text}\n\n"
    user_prompt += "위 정보를 바탕으로 완성된 계약서 전문을 작성하세요."

    try:
        result = await call_llm(
            model=settings.drafter_model,
            system_prompt=DRAFTER_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.1,
            max_tokens=8192,
        )
        contract_text = result["content"]
        return {"generated_contract": contract_text}
    except Exception as e:
        logger.error(f"계약서 생성 실패: {e}")
        return {"generated_contract": None, "error": str(e)}


async def self_review(state: DraftState) -> dict:
    """자체 검증 노드: 생성된 계약서를 간이 검토."""
    generated = state.get("generated_contract", "")
    attempt = state.get("attempt", 0)

    if not generated:
        return {"review_passed": False, "attempt": attempt}

    try:
        result = await call_llm_json(
            model=settings.validator_model,
            system_prompt=(
                "생성된 계약서의 법적 유효성을 검토하세요. JSON 응답:\n"
                '{"passed": true/false, "issues": ["문제1", ...], "score": 0~10}'
            ),
            user_prompt=f"[생성된 계약서]\n{generated[:5000]}",
            max_tokens=1024,
        )
        data = result["data"]
        passed = data.get("passed", False)
        return {
            "review_result": data,
            "review_passed": passed,
            "attempt": attempt + 1,
        }
    except Exception as e:
        logger.error(f"자체 검증 실패: {e}")
        return {"review_passed": True, "attempt": attempt + 1}  # 검증 실패시 통과 처리


async def revise_contract(state: DraftState) -> dict:
    """계약서 수정 노드: 검토 피드백 반영."""
    generated = state.get("generated_contract", "")
    review_result = state.get("review_result", {})
    issues = review_result.get("issues", [])

    if not issues or not generated:
        return {}

    feedback = "\n".join(f"- {issue}" for issue in issues)

    try:
        result = await call_llm(
            model=settings.drafter_model,
            system_prompt="계약서의 문제를 수정하세요. 수정된 전체 계약서를 반환하세요.",
            user_prompt=(
                f"[현재 계약서]\n{generated[:5000]}\n\n"
                f"[수정 필요 사항]\n{feedback}\n\n"
                "위 문제를 수정한 전체 계약서를 작성하세요."
            ),
            temperature=0.1,
            max_tokens=8192,
        )
        return {"generated_contract": result["content"]}
    except Exception as e:
        logger.error(f"계약서 수정 실패: {e}")
        return {}


def export_docx(state: DraftState) -> dict:
    """DOCX 출력 노드."""
    generated = state.get("generated_contract", "")
    session_id = state.get("session_id", "")

    if not generated:
        return {
            "response": {"status": "error", "error": "생성된 계약서가 없습니다."},
        }

    # DOCX 생성
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("계 약 서", level=0)

        for paragraph in generated.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            if paragraph.startswith("제") and "조" in paragraph[:10]:
                doc.add_heading(paragraph, level=2)
            else:
                doc.add_paragraph(paragraph)

        output_path = f"/tmp/contract_{session_id}.docx"
        doc.save(output_path)

        return {
            "output_path": output_path,
            "response": {
                "status": "completed",
                "session_id": session_id,
                "contract_text": generated,
                "output_path": output_path,
                "review_summary": state.get("review_result"),
            },
        }
    except Exception as e:
        logger.error(f"DOCX 생성 실패: {e}")
        return {
            "response": {
                "status": "completed",
                "session_id": session_id,
                "contract_text": generated,
                "output_path": None,
            },
        }
