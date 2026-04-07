"""DOCX 파서 — python-docx 기반 텍스트 추출."""

from docx import Document


def extract_text_from_docx(file_path: str) -> dict:
    """DOCX에서 텍스트를 추출한다.

    Returns:
        {
            "text": str,              # 전체 텍스트
            "paragraphs": list[str],  # 문단별 텍스트
            "tables": list[list],     # 테이블 데이터
        }
    """
    doc = Document(file_path)

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n".join(paragraphs)

    # 테이블 추출
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "text": text,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_text_from_docx_bytes(file_bytes: bytes) -> dict:
    """DOCX 바이트에서 텍스트를 추출한다."""
    import io

    doc = Document(io.BytesIO(file_bytes))

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n".join(paragraphs)

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "text": text,
        "paragraphs": paragraphs,
        "tables": tables,
    }
