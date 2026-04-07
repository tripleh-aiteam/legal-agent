"""수정 계약서 생성 및 다운로드 API."""

import logging
import tempfile
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from app.utils.db_client import fetchrow

logger = logging.getLogger(__name__)

router = APIRouter()


class RevisedContractRequest(BaseModel):
    """수정 계약서 생성 요청."""

    document_id: str
    findings: list[dict]  # [{original_text, suggested_text, title, ...}]
    output_format: str = "docx"  # "docx" | "pdf"


@router.post("/revised-contract")
async def generate_revised_contract(request: RevisedContractRequest):
    """분석 결과의 수정 제안을 반영한 계약서 파일 생성."""
    # 원본 문서 텍스트 조회
    doc = await fetchrow(
        "SELECT file_name, raw_text FROM documents WHERE id = $1",
        request.document_id,
    )
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
    if not doc["raw_text"]:
        raise HTTPException(status_code=400, detail="문서 텍스트가 없습니다.")

    raw_text = doc["raw_text"]
    file_name = doc["file_name"] or "contract"

    # 수정 제안이 있는 finding만 필터
    revisions = [
        f for f in request.findings
        if f.get("suggested_text") and f.get("original_text")
    ]

    # 원문에 수정 제안 적용
    revised_text = _apply_revisions(raw_text, revisions)

    # 출력 파일 생성
    fmt = request.output_format.lower()
    if fmt == "pdf":
        output_path = _generate_pdf(revised_text, file_name)
        media_type = "application/pdf"
        ext = "pdf"
    else:
        output_path = _generate_docx(revised_text, revisions, file_name)
        media_type = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )
        ext = "docx"

    # 다운로드 파일명 구성
    stem = Path(file_name).stem
    download_name = f"{stem}_수정본.{ext}"

    return FileResponse(
        path=output_path,
        media_type=media_type,
        filename=download_name,
    )


def _apply_revisions(raw_text: str, revisions: list[dict]) -> str:
    """원문 텍스트에 수정 제안을 적용."""
    result = raw_text
    # 긴 텍스트부터 교체 (짧은 것이 긴 것의 부분 문자열일 수 있으므로)
    sorted_revisions = sorted(
        revisions, key=lambda r: len(r["original_text"]), reverse=True,
    )

    for rev in sorted_revisions:
        original = rev["original_text"]
        suggested = rev["suggested_text"]

        if original in result:
            result = result.replace(original, suggested, 1)
        else:
            # 정확 매칭 실패 시 유사도 기반 퍼지 매칭
            best_match = _fuzzy_find(result, original)
            if best_match:
                result = result.replace(best_match, suggested, 1)
            else:
                logger.warning(
                    "수정 적용 실패 (매칭 불가): %s...",
                    original[:50],
                )

    return result


def _fuzzy_find(
    text: str, target: str, threshold: float = 0.7,
) -> str | None:
    """텍스트에서 target과 가장 유사한 부분 문자열을 찾음."""
    target_len = len(target)
    best_ratio = threshold
    best_match = None

    # 슬라이딩 윈도우로 유사 구간 탐색
    window = int(target_len * 1.3)
    for i in range(len(text) - target_len + 1):
        candidate = text[i : i + window]
        ratio = SequenceMatcher(None, target, candidate).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            # 실제 매칭 범위를 정밀하게 찾기
            blocks = SequenceMatcher(
                None, target, text[i : i + window],
            ).get_matching_blocks()
            if blocks:
                end = i + blocks[-2].b + blocks[-2].size
                best_match = text[i:end] if end > i else candidate

    return best_match


def _generate_docx(
    revised_text: str, revisions: list[dict], file_name: str,
) -> str:
    """수정 반영된 DOCX 파일 생성.

    수정된 부분을 파란색+밑줄로 표시하여 변경 사항을 시각적으로 구분.
    """
    doc = Document()

    # 문서 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # 수정된 텍스트 목록 (하이라이트용)
    suggested_texts = {
        r["suggested_text"] for r in revisions if r.get("suggested_text")
    }

    for line in revised_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # 조항 제목 감지
        if line.startswith("제") and "조" in line[:10]:
            doc.add_heading(line, level=2)
            continue

        para = doc.add_paragraph()

        # 수정된 부분이 포함된 라인이면 하이라이트
        highlighted = False
        for suggested in suggested_texts:
            if suggested in line:
                # 수정된 부분 전후로 분리하여 색상 표시
                parts = line.split(suggested, 1)
                if parts[0]:
                    run = para.add_run(parts[0])
                    run.font.size = Pt(10)
                # 수정된 부분: 파란색 + 밑줄
                run = para.add_run(suggested)
                run.font.color.rgb = RGBColor(0, 0, 180)
                run.font.underline = True
                run.font.size = Pt(10)
                if len(parts) > 1 and parts[1]:
                    run = para.add_run(parts[1])
                    run.font.size = Pt(10)
                highlighted = True
                break

        if not highlighted:
            run = para.add_run(line)
            run.font.size = Pt(10)

    # 수정 요약 페이지 추가
    if revisions:
        doc.add_page_break()
        doc.add_heading("수정 사항 요약", level=1)
        for i, rev in enumerate(revisions, 1):
            doc.add_heading(
                f"{i}. {rev.get('title', '수정 사항')}", level=3,
            )

            # 원문
            p = doc.add_paragraph()
            run = p.add_run("원문: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(rev["original_text"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 0, 0)

            # 수정안
            p = doc.add_paragraph()
            run = p.add_run("수정: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(rev["suggested_text"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 180)

            # 수정 이유
            reason = rev.get("suggestion_reason")
            if reason:
                p = doc.add_paragraph()
                run = p.add_run("사유: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(reason)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)

    # 임시 파일로 저장
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, prefix="revised_",
    )
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _generate_pdf(revised_text: str, file_name: str) -> str:
    """수정 반영된 PDF 파일 생성.

    python-docx로 DOCX 먼저 생성 후 pymupdf로 텍스트 기반 PDF 생성.
    """
    import fitz  # pymupdf

    pdf = fitz.open()
    # A4 사이즈
    width, height = 595, 842
    margin = 50
    y = margin
    line_height = 14
    font_size = 10

    page = pdf.new_page(width=width, height=height)

    for line in revised_text.split("\n"):
        line = line.strip()
        if not line:
            y += line_height * 0.5
            continue

        # 페이지 넘김
        if y + line_height > height - margin:
            page = pdf.new_page(width=width, height=height)
            y = margin

        # 조항 제목은 볼드+크게
        if line.startswith("제") and "조" in line[:10]:
            page.insert_text(
                (margin, y), line,
                fontsize=font_size + 2,
                fontname="helv",
            )
            y += line_height * 1.5
        else:
            # 긴 줄 자동 줄바꿈
            max_chars = 55  # 한 줄 최대 글자수 (한글 기준)
            while len(line) > max_chars:
                page.insert_text(
                    (margin, y), line[:max_chars],
                    fontsize=font_size,
                    fontname="helv",
                )
                line = line[max_chars:]
                y += line_height
                if y + line_height > height - margin:
                    page = pdf.new_page(width=width, height=height)
                    y = margin

            if line:
                page.insert_text(
                    (margin, y), line,
                    fontsize=font_size,
                    fontname="helv",
                )
                y += line_height

    tmp = tempfile.NamedTemporaryFile(
        suffix=".pdf", delete=False, prefix="revised_",
    )
    pdf.save(tmp.name)
    pdf.close()
    tmp.close()
    return tmp.name
